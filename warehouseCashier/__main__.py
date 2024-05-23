from docx import Document

from openpyxl import Workbook

from openpyxl.worksheet.table import Table, TableStyleInfo

from openpyxl.styles import Font


def outputWord(totalHarga, kumpulanBarang) :
    # Buat dokumen Word baru
    document = Document()

    # Tambahkan teks output ke dokumen
    output_text = "=================================================\n" \
                "Terima Kasih Telah Berbelanja!" \
        
    document.add_paragraph(output_text)
    
    for key, value in kumpulanBarang.items():
        barangText = f"Nama Barang : {key}  Jumlah :  {value[0]}  Harga :  {value[1]}"
        document.add_paragraph(barangText)
    
    
    output_text2 = f"Total Harga : {totalHarga} \n=================================================\n" 
    
    document.add_paragraph(output_text2)
    
    # Simpan dokumen sebagai Word
    document.save('output.docx')

def outputExcel(gudang) :
    wb = Workbook()
    ws = wb.active
    
    #nama column
    treeData = [["ID Barang", "Nama Barang", "Stok", "Harga"]]
    for rowTree in treeData :
        ws.append(rowTree)
    #add barang per row
    for key, value in gudang.items():
        row = [f"{key}", f"{value[0]}", f"{value[1]}", f"{value[2]}"]
        ws.append(row)
    #font
    ft = Font(bold=True)
    for row in ws["A1:C1"] :
        for cell in row:
            cell.font = ft

    wb.save("gudang.xlsx")

    

def outputStruk(kumpulanBarang) :
    for key, value in kumpulanBarang.items():
        print('Nama Barang :', format(key), ' Jumlah : ', format(value[0]) ,' Harga : ',format(value[1]))
        return
def crud(gudang) :
    while True :
        print("Menu Manajemen Gudang :")
        print("1. Menambah Barang Baru")
        print("2. Edit Barang")
        print("3. Menghapus Barang")
        print("4. Menampilkan Gudang & export ke excel")
        print("5. Kembali ke Menu Utama")
        pilihanMenugudang = int(input("Masukan pilihan menu (1/2/3/4) : "))
        
        if pilihanMenugudang == 1 :
            for key, value in gudang.items():
                print("ID :",format(key),'Nama Barang :', format(value[0]), ' stok : ', format(value[1]) ,' Harga : ',format(value[2]))
            keySet = str(input("Masukan Key Baru (format: ID-nomor) :"))
            nameSet = str(input("Masukan Nama Barang (str) : "))
            stokSet = int(input("Masukan stok barang (int) : "))
            priceSet = int(input("Masukan harga barang (int) : "))
            gudang[keySet] = [nameSet, stokSet, priceSet]
            for key, value in gudang.items():
                print("ID :",format(key),'Nama Barang :', format(value[0]), ' stok : ', format(value[1]) ,' Harga : ',format(value[2]))
        
        elif pilihanMenugudang == 2 :
            for key, value in gudang.items():
                print("ID :",format(key),'Nama Barang :', format(value[0]), ' stok : ', format(value[1]) ,' Harga : ',format(value[2]))
            keySet = str(input("Masukan Key yang mau diedit (format: ID-nomor) :"))
            nameSet = str(input("Masukan Nama Barang (str) : "))
            stokSet = int(input("Masukan stok barang (int) : "))
            priceSet = int(input("Masukan harga barang (int) : "))
            gudang[keySet] = [nameSet, stokSet, priceSet]
            for key, value in gudang.items():
                print("ID :",format(key),'Nama Barang :', format(value[0]), ' stok : ', format(value[1]) ,' Harga : ',format(value[2]))
                
        elif pilihanMenugudang == 3 :
            for key, value in gudang.items():
                print("ID :",format(key),'Nama Barang :', format(value[0]), ' stok : ', format(value[1]) ,' Harga : ',format(value[2]))
            keySet = str(input("Masukan Key yang mau dihapus (format: ID-nomor) :"))
            del gudang[keySet]
            for key, value in gudang.items():
                print("ID :",format(key),'Nama Barang :', format(value[0]), ' stok : ', format(value[1]) ,' Harga : ',format(value[2]))
        elif pilihanMenugudang == 4 :
            for key, value in gudang.items():
                print("ID :",format(key),'Nama Barang :', format(value[0]), ' stok : ', format(value[1]) ,' Harga : ',format(value[2]))
            pilihanExcel = str(input("Apakah ingin export ke excel? (Yes/No) : "))
            if pilihanExcel.lower() == "yes" :
                outputExcel(gudang)
                print("File Excel telah terbuat!")
            elif pilihanExcel.lower() == "no" :
                continue
        elif pilihanMenugudang == 5 :
            main(gudang)
            
def kasir(gudang) :
    kumpulanBarang = {} # namaBarang : [jumlah barang, jumlah harga ]
    kumpulanHarga = []
    x = 1
    while True :
        for key, value in gudang.items():
                print("ID :",format(key),'Nama Barang :', format(value[0]), ' stok : ', format(value[1]) ,' Harga : ',format(value[2]))
        pilihanBarang = str(input("Masukan ID barang : "))
        jumlahBarang = int(input("Masukan jumlah barang tersebut yang ingin dibeli (int) : "))
        
        ambilValueharga = gudang[pilihanBarang][2]
        ambilNamabarang = gudang[pilihanBarang][0]
        gudang[pilihanBarang][1] -= jumlahBarang # kurangi stok
        
        hitungHargabarang = ambilValueharga * jumlahBarang #barang * jumlah beli
        
        kumpulanBarang[ambilNamabarang+f"| Barang ke {x} "] = [jumlahBarang, hitungHargabarang] #tambah ke dict
        kumpulanHarga.append(hitungHargabarang)
        
        x += 1
            
        pilihanUser = str(input("Apakah ingin menambah barang ? (Yes/No) : "))
        if pilihanUser.lower() == "yes" :
            continue
        elif pilihanUser.lower() == "no" :
            totalHarga = sum(kumpulanHarga)
            print("=================================================")
            print("Terima Kasih Telah Berbelanja!")
            for key, value in kumpulanBarang.items():
                print('Nama Barang :', format(key), ' Jumlah : ', format(value[0]) ,' Harga : ',format(value[1]))
            print(f"Total Harga : {totalHarga}")
            print("=================================================")
            # varforPrint2 = print(f"{varforPrint}")
            outputWord(totalHarga, kumpulanBarang) 
            kumpulanHarga.clear()
            kumpulanBarang.clear()
            x = 1
            main(gudang)
            
def main(gudang) :
    while True :
        print("Pilihan Menu :")
        print("1. Manajemen Gudang")
        print("2. Kasir")
        print("3. Keluar")
        pilihanMenu = int(input("Masukan pilihan (1/2/3) : "))
        
        if pilihanMenu == 1 :
            crud(gudang)
        elif pilihanMenu == 2 :
            kasir(gudang)
        elif pilihanMenu == 3 :
            break
gudang = {
    "ID-1" : ["Rinso", 500, 6000],
    "ID-2" : ["Sosis", 450, 1500],
    "ID-3" : ["Citrun", 350, 500],
    "ID-4" : ["Susu Kotak", 280, 2000],
    "ID-5" : ["Oreo", 330, 2000],
}

main(gudang)
        